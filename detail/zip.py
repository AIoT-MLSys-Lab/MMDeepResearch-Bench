# inspect_daily_docx.py
# 用法示例：
#   python inspect_daily_docx.py --docx "E:\path\to\daily.docx" --extract --out_dir ".\_docx_media" --max_text 80
# 如果你的项目里有 question_parser.py，还可以加：
#   python inspect_daily_docx.py --docx "..." --use_question_parser

from __future__ import annotations

import argparse
import os
import sys
import zipfile
from pathlib import Path
from datetime import datetime
import xml.etree.ElementTree as ET

def human_bytes(n: int) -> str:
    for u in ["B", "KB", "MB", "GB"]:
        if n < 1024:
            return f"{n:.0f}{u}"
        n /= 1024
    return f"{n:.1f}TB"

def inspect_zip_media(docx: Path, extract: bool, out_dir: Path) -> None:
    print("\n=== ZIP / Embedded Media Check ===")
    with zipfile.ZipFile(docx, "r") as z:
        names = z.namelist()

        media = [n for n in names if n.startswith("word/media/")]
        print(f"embedded media files under word/media/: {len(media)}")

        if media:
            for n in media[:30]:
                info = z.getinfo(n)
                print(f"  - {n} ({human_bytes(info.file_size)})")
            if len(media) > 30:
                print(f"  ... ({len(media)-30} more)")

        if extract and media:
            out_dir.mkdir(parents=True, exist_ok=True)
            for n in media:
                target = out_dir / Path(n).name
                with z.open(n) as src, open(target, "wb") as dst:
                    dst.write(src.read())
            print(f"extracted {len(media)} media files -> {out_dir.resolve()}")

def inspect_relationships(docx: Path) -> None:
    print("\n=== Relationships (image embed / external) ===")
    with zipfile.ZipFile(docx, "r") as z:
        rels_files = [n for n in z.namelist() if n.startswith("word/_rels/") and n.endswith(".rels")]
        if not rels_files:
            print("no .rels found under word/_rels/")
            return

        ns = {"r": "http://schemas.openxmlformats.org/package/2006/relationships"}
        total_img_rels = 0
        external_img_rels = 0

        for rf in sorted(rels_files):
            data = z.read(rf)
            try:
                root = ET.fromstring(data)
            except Exception as e:
                print(f"[warn] failed to parse {rf}: {e}")
                continue

            img_rels = []
            for rel in root.findall("r:Relationship", ns):
                rtype = (rel.attrib.get("Type") or "")
                if rtype.endswith("/image"):
                    img_rels.append(rel)

            if not img_rels:
                continue

            print(f"\n{rf}: image relationships = {len(img_rels)}")
            for rel in img_rels[:40]:
                rid = rel.attrib.get("Id")
                tgt = rel.attrib.get("Target")
                mode = rel.attrib.get("TargetMode", "Internal")
                print(f"  - {rid}: TargetMode={mode} Target={tgt}")
                total_img_rels += 1
                if mode.lower() == "external":
                    external_img_rels += 1

            if len(img_rels) > 40:
                print(f"  ... ({len(img_rels)-40} more in this rels file)")

        print(f"\nsummary: image rels={total_img_rels}, external_image_rels={external_img_rels}")

def inspect_text(docx: Path, max_text: int) -> None:
    print("\n=== python-docx Text Preview ===")
    try:
        from docx import Document  # python-docx
    except Exception as e:
        print(f"[warn] python-docx not available: {e}")
        return

    try:
        d = Document(str(docx))
    except Exception as e:
        print(f"[warn] failed to open with python-docx: {e}")
        return

    paras = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
    print(f"non-empty paragraphs: {len(paras)}")
    for i, t in enumerate(paras[:max_text], 1):
        print(f"{i:03d}: {t}")

    try:
        print(f"tables: {len(d.tables)}")
    except Exception:
        pass

def try_question_parser(docx: Path) -> None:
    print("\n=== Your Project's question_parser.load_questions_from_docx (optional) ===")
    try:
        from .question_parser import load_questions_from_docx  # type: ignore
    except Exception as e:
        print(f"[skip] cannot import question_parser.load_questions_from_docx: {e}")
        return

    try:
        qs = load_questions_from_docx(docx)
    except Exception as e:
        print(f"[warn] load_questions_from_docx crashed: {e}")
        return

    print(f"loaded questions: {len(qs)}")
    for q in qs[:30]:
        qid = getattr(q, "qid", "")
        title = getattr(q, "title", "")
        text = getattr(q, "text", "")
        print(f"  - {qid} | {title} | text_len={len(text or '')}")
    if len(qs) > 30:
        print(f"  ... ({len(qs)-30} more)")

def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", required=True, help="path to daily docx")
    ap.add_argument("--extract", action="store_true", help="extract embedded media to out_dir")
    ap.add_argument("--out_dir", default="./_docx_media", help="where to extract media")
    ap.add_argument("--max_text", type=int, default=60, help="max non-empty paragraphs to print")
    ap.add_argument("--use_question_parser", action="store_true", help="also run your repo's load_questions_from_docx if importable")
    args = ap.parse_args()

    docx = Path(args.docx).expanduser().resolve()
    print("DOCX:", docx)
    print("exists:", docx.exists())
    if not docx.exists():
        sys.exit(1)

    st = docx.stat()
    print("size:", human_bytes(st.st_size))
    print("mtime:", datetime.fromtimestamp(st.st_mtime).isoformat(sep=" ", timespec="seconds"))

    inspect_zip_media(docx, extract=args.extract, out_dir=Path(args.out_dir))
    inspect_relationships(docx)
    inspect_text(docx, max_text=args.max_text)

    if args.use_question_parser:
        try_question_parser(docx)

    print("\nNOTE:")
    print("  - 如果 embedded media 数量 > 0，说明 Word 里确实嵌入了图片（word/media/ 下）。")
    print("  - 如果 image relationships 里出现 TargetMode=External，说明图片可能是外链/链接图，不一定会打包进 docx。")

if __name__ == "__main__":
    main()
